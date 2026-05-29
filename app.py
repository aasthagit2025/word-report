import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io

def set_cell_background(cell, color_hex):
    """Applies a background fill color to a cell block."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def process_survey_to_docx(csv_file):
    df = pd.read_csv(csv_file)
    doc = Document()
    
    # Page setup: Standard margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Exact Header Construction
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    r1 = title_p.add_run("BWH Hotels\n")
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r1.font.bold = True
    
    r2 = title_p.add_run("2025 Member Survey,Open End Comments\n")
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    
    r3 = title_p.add_run("Surestay and Collections,North America,,,")
    r3.font.name = 'Arial'
    r3.font.size = Pt(11)
    
    r4 = title_p.add_run("*SureStay Responses Highlighted in Blue")
    r4.font.name = 'Arial'
    r4.font.size = Pt(11)
    r4.font.italic = True
    
    # 2. Match Target Questions
    q7_columns = [col for col in df.columns if col.startswith("Q7 -")]
    q8_column = "Q8 - Is there anything else you would like us to know? Any additional thoughts you would like to share are appreciated."
    target_questions = q7_columns + ([q8_column] if q8_column in df.columns else [])
    
    # 3. Build text items sequentially without interactive question prompts
    for q_full_name in target_questions:
        valid_responses = df[[q_full_name, 'Brand']].dropna(subset=[q_full_name])
        valid_responses = valid_responses[valid_responses[q_full_name].astype(str).str.strip() != ""]
        
        if valid_responses.empty:
            continue
            
        doc.add_paragraph("") # Space line
        
        # Split the question string for dual-line formatting if it contains a sub-header
        if " - " in q_full_name:
            parts = q_full_name.split(" - ", 1)
            q_header = parts[0] + " - " + parts[1]
            
            # If it has a sub-component (like Education & Training...) split it neatly
            if "Dept’s" in q_header or "process" in q_header:
                main_q = parts[0] + " - " + parts[1].split(".")[0]
                doc.add_paragraph().add_run(main_q).font.bold = True
                
                # Check for sub-text
                sub_parts = parts[1].split(".", 1)
                if len(sub_parts) > 1 and sub_parts[1].strip():
                    doc.add_paragraph().add_run(sub_parts[1].strip())
            else:
                q_p = doc.add_paragraph()
                qr = q_p.add_run(q_full_name)
                qr.font.name = 'Arial'
                qr.font.bold = True
        else:
            q_p = doc.add_paragraph()
            qr = q_p.add_run(q_full_name)
            qr.font.name = 'Arial'
            qr.font.bold = True

        # Render rows as plain text lines inside a seamless bordered cell container
        table = doc.add_table(rows=0, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False
        
        # Apply minimalist light outer borders matching the reference layout
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:left w:val="none"/>'
            '<w:right w:val="none"/>'
            '<w:insideH w:val="none"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        for _, row in valid_responses.iterrows():
            text_comment = str(row[q_full_name]).strip()
            brand_type = str(row['Brand']).upper()
            
            row_cells = table.add_row().cells
            cell = row_cells[0]
            cell.width = Inches(6.5)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            
            text_run = p.add_run(text_comment)
            text_run.font.name = 'Arial'
            text_run.font.size = Pt(10.5)
            
            # Apply background tint color if it belongs to a SureStay property code
            if brand_type.startswith("SS"):
                set_cell_background(cell, "DCE6F1") # Soft Blue Hex
                
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# --- Streamlit UI Setup ---
st.set_page_config(page_title="Survey Format Pipeline", page_icon="📋")
st.title("📋 Open-End Response Alignment Tool")

uploaded_file = st.file_uploader("Upload Survey CSV Data Source", type=["csv"])

if uploaded_file is not None:
    if st.button("Convert to Reference Layout File", type="primary"):
        with st.spinner("Processing text streams..."):
            try:
                output_buffer = process_survey_to_docx(uploaded_file)
                st.success("Format processing complete!")
                st.download_button(
                    label="📥 Download Exact Match Document",
                    data=output_buffer,
                    file_name="BWH_Survey_Open_End_Comments.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Execution Error: {e}")